"""Generate management- and audit-ready Word data quality reports."""

from __future__ import annotations

from collections import Counter
from datetime import datetime
from io import BytesIO
from typing import Any, Iterable, Mapping, Sequence
from uuid import uuid4

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from src.analyzer import (
    detect_business_column_type,
    friendly_column_name,
)
from src.data_quality import (
    CleaningAuditEntry,
    audit_records,
    business_action_text,
    classify_missing_values,
    missing_status_by_column,
)
from src.integrity import integrity_issues_frame, validate_integrity
from src.logger_setup import setup_logger

logger = setup_logger(__name__)

PREVIEW_ROWS = 8
PREVIEW_COLUMNS = 8
MAX_OUTLIER_ROWS = 20
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_TOP_BOTTOM = 80
CELL_MARGIN_START_END = 120

NAVY = RGBColor(15, 23, 42)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(71, 85, 105)
WHITE = RGBColor(255, 255, 255)
LIGHT_FILL = "F2F4F7"
BLUE_FILL = "E8EEF5"
WARNING_FILL = "FFF7ED"
INFO_FILL = "EFF6FF"
BORDER = "CBD5E1"

SECTION_TITLES = {
    "overview": "Dataset Overview",
    "quality": "Data Quality and Cleaning Summary",
    "statistics": "Column Statistics",
    "outliers": "Values Worth Reviewing",
    "charts": "Charts",
    "preview": "Data Preview",
    "methodology": "Methodology and Limitations",
}
SECTION_ORDER = (
    "overview",
    "quality",
    "statistics",
    "outliers",
    "charts",
    "preview",
    "methodology",
)
DEFAULT_SECTIONS = (
    "overview",
    "quality",
    "statistics",
    "outliers",
    "methodology",
)


class ReportError(Exception):
    """Raised when report generation fails."""


def _set_run_font(
    run: Any,
    *,
    name: str = "Calibri",
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _configure_document(doc: Document) -> None:
    """Apply the standard-business-brief preset and page geometry."""
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = NAVY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def _field_run(paragraph: Any, instruction: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    field = OxmlElement("w:instrText")
    field.set(qn("xml:space"), "preserve")
    field.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.extend([begin, field, separate, text, end])
    _set_run_font(run, size=8.5, color=MUTED)


def _add_page_footer(doc: Document, generated_at: datetime) -> None:
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(
        f"Data Quality Report  |  {generated_at:%Y-%m-%d}  |  Page "
    )
    _set_run_font(run, size=8.5, color=MUTED)
    _field_run(paragraph, "PAGE")
    run = paragraph.add_run(" of ")
    _set_run_font(run, size=8.5, color=MUTED)
    _field_run(paragraph, "NUMPAGES")


def _add_running_header(doc: Document, report_id: str) -> None:
    header = doc.sections[0].header
    header.is_linked_to_previous = False
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run(f"Excel Automation Toolkit  |  Report {report_id}")
    _set_run_font(run, size=8.5, color=MUTED)


def _display(value: Any, *, decimals: int = 2) -> str:
    if value is None:
        return "—"
    try:
        if bool(pd.isna(value)):
            return "—"
    except (TypeError, ValueError):
        pass
    if isinstance(value, (pd.Timestamp, datetime)):
        return pd.Timestamp(value).strftime("%Y-%m-%d")
    if isinstance(value, float):
        if value.is_integer():
            return f"{int(value):,}"
        return f"{value:,.{decimals}f}"
    if isinstance(value, int):
        return f"{value:,}"
    return str(value)


def _set_cell_margins(cell: Any) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (
        ("top", CELL_MARGIN_TOP_BOTTOM),
        ("bottom", CELL_MARGIN_TOP_BOTTOM),
        ("start", CELL_MARGIN_START_END),
        ("end", CELL_MARGIN_START_END),
    ):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table: Any, widths_dxa: Sequence[int]) -> None:
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError("Word table column widths must total 9360 DXA.")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.get_or_add_tcW()
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _fit_widths(column_names: Sequence[str]) -> list[int]:
    """Return content-aware widths that sum to the preset content width."""
    weights = []
    for name in column_names:
        lower = name.lower()
        if any(token in lower for token in ("reason", "explanation", "action")):
            weights.append(2.4)
        elif any(token in lower for token in ("source", "description", "scope")):
            weights.append(1.6)
        elif any(token in lower for token in ("count", "blank", "value", "date")):
            weights.append(1.0)
        else:
            weights.append(1.25)
    total = sum(weights) or 1
    widths = [int(CONTENT_WIDTH_DXA * weight / total) for weight in weights]
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def _add_table(
    doc: Document,
    headers: Sequence[str],
    rows: Iterable[Sequence[Any]],
    *,
    widths_dxa: Sequence[int] | None = None,
    header_fill: str = LIGHT_FILL,
    keep_together: bool = False,
) -> Any:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = str(header)
        cell._tc.get_or_add_tcPr().append(_cell_shading(header_fill))
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                _set_run_font(run, size=8.5, color=NAVY, bold=True)

    for row_values in rows:
        row = table.add_row()
        for index, value in enumerate(row_values):
            cell = row.cells[index]
            cell.text = _display(value)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    _set_run_font(run, size=8.5, color=NAVY)
                if isinstance(value, (int, float)) and not isinstance(value, bool):
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for row_index, row in enumerate(table.rows):
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        if keep_together and row_index < len(table.rows) - 1:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.keep_with_next = True

    _set_table_geometry(table, list(widths_dxa or _fit_widths(headers)))
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    return table


def _cell_shading(fill: str) -> OxmlElement:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    return shading


def _shade_paragraph(paragraph: Any, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = p_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        p_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _add_note(doc: Document, text: str, *, warning: bool = False) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.12)
    _shade_paragraph(paragraph, WARNING_FILL if warning else INFO_FILL)
    label = paragraph.add_run("Review note: " if warning else "Note: ")
    _set_run_font(label, size=10, color=NAVY, bold=True)
    run = paragraph.add_run(text)
    _set_run_font(run, size=10, color=NAVY)


def _add_title_area(
    doc: Document,
    *,
    generated_at: datetime,
    dataset_name: str,
    source_files: list[str],
    report_id: str,
) -> None:
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(8)
    kicker.paragraph_format.space_after = Pt(4)
    run = kicker.add_run("DATA GOVERNANCE AND QUALITY")
    _set_run_font(run, size=9, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("Data Quality Report")
    _set_run_font(run, size=26, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run(dataset_name)
    _set_run_font(run, size=13, color=MUTED)

    metadata = [
        ("Generated", f"{generated_at:%Y-%m-%d %H:%M}"),
        ("Source files", ", ".join(source_files) if source_files else "—"),
        ("Toolkit", "Excel Automation Toolkit"),
        ("Report ID", report_id),
    ]
    _add_table(
        doc,
        ["Report detail", "Value"],
        metadata,
        widths_dxa=[2400, 6960],
        header_fill=BLUE_FILL,
    )


def _executive_summary(
    df: pd.DataFrame,
    source_files: list[str],
    cleaning_report: Mapping[str, Any],
    outlier_df: pd.DataFrame | None,
) -> str:
    duplicates = int(cleaning_report.get("duplicates_removed", 0))
    reviewed = int(cleaning_report.get("missing_values_reviewed", 0))
    structural = int(cleaning_report.get(
        "unavailable_from_source",
        cleaning_report.get("structural_missing_after", 0),
    ))
    approved_blank = int(cleaning_report.get("approved_unchanged", 0))
    deterministic = int(cleaning_report.get("deterministic_recoveries", 0))
    outliers = 0 if outlier_df is None else len(outlier_df)
    integrity_passed = bool(cleaning_report.get("integrity_passed", True))
    number_words = {
        0: "zero", 1: "one", 2: "two", 3: "three", 4: "four",
        5: "five", 6: "six", 7: "seven", 8: "eight", 9: "nine",
        10: "ten", 11: "eleven", 12: "twelve", 13: "thirteen",
        14: "fourteen", 15: "fifteen", 16: "sixteen",
        17: "seventeen", 18: "eighteen", 19: "nineteen", 20: "twenty",
    }
    word = lambda value: number_words.get(value, f"{value:,}")
    return (
        f"The approved dataset contains {len(df):,} records from "
        f"{word(len(source_files))} source "
        f"{'file' if len(source_files) == 1 else 'files'}. "
        f"{word(duplicates).capitalize()} repeated "
        f"{'record was' if duplicates == 1 else 'records were'} removed. "
        f"{word(reviewed).capitalize()} missing "
        f"{'value was' if reviewed == 1 else 'values were'} reviewed: "
        f"{word(deterministic)} "
        f"{'quantity was' if deterministic == 1 else 'quantities were'} "
        "recovered from a validated arithmetic relationship, and "
        f"{word(approved_blank)} "
        f"{'value was' if approved_blank == 1 else 'values were'} approved "
        "to remain blank. A further "
        f"{structural:,} "
        f"{'cell was' if structural == 1 else 'cells were'} unavailable because "
        "their source files did not contain the corresponding fields. "
        f"{word(outliers).capitalize()} unusual "
        f"{'value was' if outliers == 1 else 'values were'} flagged for review. "
        "Relationship validation "
        f"{'passed' if integrity_passed else 'failed'}."
    )


def _add_overview(
    doc: Document,
    df: pd.DataFrame,
    source_files: list[str],
    source_schemas: Mapping[str, Iterable[str]] | None,
) -> None:
    overview_rows: list[tuple[str, object]] = [
        ("Records", len(df)),
        ("Columns", len(df.columns)),
        ("Source files", len(source_files)),
    ]
    date_columns = [
        column
        for column in df.columns
        if detect_business_column_type(df[column]) == "date"
    ]
    if date_columns:
        parsed = pd.to_datetime(df[date_columns[0]], errors="coerce").dropna()
        if not parsed.empty:
            overview_rows.extend([
                ("Date range begins", parsed.min().strftime("%Y-%m-%d")),
                ("Date range ends", parsed.max().strftime("%Y-%m-%d")),
            ])
    _add_table(
        doc,
        ["Metric", "Value"],
        overview_rows,
        widths_dxa=[3000, 6360],
    )

    if "source_file" in df.columns:
        contributions = (
            df["source_file"].value_counts(dropna=False).rename_axis("Source")
        )
        _add_table(
            doc,
            ["Source File", "Records"],
            [(source, int(count)) for source, count in contributions.items()],
            widths_dxa=[6500, 2860],
            keep_together=True,
        )

    if source_schemas:
        union = set().union(*(set(columns) for columns in source_schemas.values()))
        schema_rows = []
        for source in source_files:
            available = set(source_schemas.get(source, []))
            unavailable = sorted(union - available)
            schema_rows.append((
                source,
                ", ".join(friendly_column_name(column) for column in unavailable)
                if unavailable
                else "All combined columns supplied",
            ))
        _add_table(
            doc,
            ["Source File", "Columns Not Provided"],
            schema_rows,
            widths_dxa=[3300, 6060],
        )


def _group_audit_actions(
    audit: Iterable[CleaningAuditEntry | Mapping[str, Any]],
    *,
    action_types: set[str] | None = None,
) -> list[tuple[str, str, str, int, str]]:
    records = audit_records(audit)
    grouped: Counter[tuple[str, str, str, str]] = Counter()
    for record in records:
        if record.get("action_type") == "structural_blank":
            continue
        if (
            action_types is not None
            and record.get("action_type") not in action_types
        ):
            continue
        key = (
            business_action_text(record.get("action", "Approved action")),
            friendly_column_name(record.get("column", "—")),
            str(record.get("strategy_scope", "—")),
            _display(record.get("resulting_value")),
        )
        grouped[key] += max(int(record.get("rows_removed", 0)), 1)
    return [
        (action, column, scope, affected, replacement)
        for (action, column, scope, replacement), affected in grouped.items()
    ]


def _add_quality_summary(
    doc: Document,
    cleaning_report: Mapping[str, Any],
    audit: Iterable[CleaningAuditEntry | Mapping[str, Any]],
) -> None:
    rows = [
        ("Rows before cleaning", cleaning_report.get("rows_before", "—")),
        ("Rows after cleaning", cleaning_report.get("rows_after", "—")),
        ("Repeated rows removed", cleaning_report.get("duplicates_removed", 0)),
        (
            "Incomplete rows removed",
            cleaning_report.get("incomplete_rows_removed", 0),
        ),
        (
            "Missing values reviewed",
            cleaning_report.get("missing_values_reviewed", 0),
        ),
        (
            "Values recovered or filled",
            cleaning_report.get("values_changed", cleaning_report.get("values_filled", 0)),
        ),
        (
            "Approved to remain blank",
            cleaning_report.get("approved_unchanged", 0),
        ),
        (
            "Decisions pending",
            cleaning_report.get("decisions_pending", 0),
        ),
        (
            "Unavailable from source",
            cleaning_report.get(
                "unavailable_from_source",
                cleaning_report.get("structural_missing_after", 0),
            ),
        ),
        (
            "Integrity failures",
            cleaning_report.get(
                "integrity_failures",
                cleaning_report.get("severe_integrity_issue_count", 0),
            ),
        ),
    ]
    _add_table(doc, ["Quality Measure", "Count"], rows, widths_dxa=[6500, 2860])

    deterministic_actions = _group_audit_actions(
        audit,
        action_types={"deterministic_recovery"},
    )
    estimated_actions = _group_audit_actions(
        audit,
        action_types={"fill_blank"},
    )
    other_actions = _group_audit_actions(
        audit,
        action_types={"remove_duplicate", "remove_incomplete_row"},
    )
    if deterministic_actions:
        doc.add_heading("Deterministic Recoveries", level=2)
        deterministic_rows = []
        for record in audit_records(audit):
            if record.get("action_type") != "deterministic_recovery":
                continue
            action = str(record.get("action", ""))
            action = action.replace("total", "Total").replace(
                "unit_price", "Unit Price"
            )
            formula = str(record.get("formula_or_strategy", "—"))
            formula = formula.replace("total", "Total").replace(
                "unit_price", "Unit Price"
            ).replace("quantity", "Quantity")
            deterministic_rows.append([
                record.get("business_record_identifier")
                or record.get("row_identifier", "—"),
                record.get("source_file", "—"),
                friendly_column_name(record.get("column", "—")),
                record.get("original_state", "Blank"),
                _display(record.get("resulting_value")),
                formula,
            ])
        _add_table(
            doc,
            ["Record ID", "Source File", "Field", "Original State", "Result", "Formula"],
            deterministic_rows,
            widths_dxa=[1100, 1550, 950, 1850, 800, 3110],
        )
    if estimated_actions:
        doc.add_heading("Estimated Replacements", level=2)
        _add_table(
            doc,
            ["Action", "Column", "Calculation Scope", "Affected", "Replacement"],
            estimated_actions,
            widths_dxa=[2100, 1500, 2500, 1100, 2160],
        )
    if other_actions:
        doc.add_heading("Approved Cleaning Actions", level=2)
        _add_table(
            doc,
            ["Action", "Column", "Calculation Scope", "Affected", "Replacement"],
            other_actions,
            widths_dxa=[2100, 1500, 2500, 1100, 2160],
        )
    if (
        not deterministic_actions
        and not estimated_actions
        and not other_actions
        and cleaning_report.get("missing_actions")
    ):
        doc.add_heading("Approved Cleaning Actions", level=2)
        for action in cleaning_report.get("missing_actions", []):
            paragraph = doc.add_paragraph(style="List Bullet")
            run = paragraph.add_run(business_action_text(action))
            _set_run_font(run, size=11, color=NAVY)
    elif not deterministic_actions and not estimated_actions and not other_actions:
        doc.add_paragraph("No missing-value replacements or row removals were recorded.")

    doc.add_heading("Integrity Checks", level=2)
    passed = bool(cleaning_report.get("integrity_passed", True))
    _add_table(
        doc,
        ["Check", "Result", "Issues"],
        [[
            "Validated arithmetic relationships",
            "Passed" if passed else "Failed",
            cleaning_report.get("severe_integrity_issue_count", 0),
        ]],
        widths_dxa=[5000, 2200, 2160],
    )
def _business_statistics(
    df: pd.DataFrame,
    source_schemas: Mapping[str, Iterable[str]] | None,
    cleaning_audit: Iterable[CleaningAuditEntry | Mapping[str, Any]],
) -> dict[str, pd.DataFrame]:
    records: dict[str, list[dict[str, Any]]] = {
        "numeric": [],
        "categorical": [],
        "date": [],
        "identifier": [],
        "empty": [],
    }
    statuses = missing_status_by_column(
        df,
        source_schemas,
        cleaning_audit,
    )
    for column in df.columns:
        series = df[column]
        kind = detect_business_column_type(series)
        non_null = series.dropna()
        status = statuses.get(column)
        blank_status = {
            "Approved blank": status.approved_blank if status else 0,
            "Unavailable from source": (
                status.unavailable_from_source if status else 0
            ),
            "Decisions pending": status.decisions_pending if status else 0,
        }
        base = {
            "Column": friendly_column_name(column),
        }
        if kind == "numeric":
            records[kind].append({
                **base,
                "Non-Missing": int(series.notna().sum()),
                "Average": non_null.mean() if not non_null.empty else None,
                "Median": non_null.median() if not non_null.empty else None,
                "Minimum": non_null.min() if not non_null.empty else None,
                "Maximum": non_null.max() if not non_null.empty else None,
                **blank_status,
            })
        elif kind == "categorical":
            mode = non_null.mode(dropna=True)
            most_common = mode.iloc[0] if not mode.empty else None
            records[kind].append({
                **base,
                "Unique Values": int(non_null.nunique()),
                "Most Common": most_common,
                "Frequency": int((non_null == most_common).sum())
                if most_common is not None
                else 0,
                **blank_status,
            })
        elif kind == "date":
            parsed = pd.to_datetime(non_null, errors="coerce").dropna()
            records[kind].append({
                **base,
                "Earliest": parsed.min() if not parsed.empty else None,
                "Latest": parsed.max() if not parsed.empty else None,
                "Distinct Dates": int(parsed.dt.date.nunique())
                if not parsed.empty
                else 0,
                **blank_status,
            })
        elif kind == "identifier":
            records[kind].append({
                "Column": friendly_column_name(column),
                "Total Records": len(series),
                "Unique Values": int(non_null.nunique()),
                "Repeated Identifiers": int(non_null.duplicated(keep=False).sum()),
                **blank_status,
            })
        else:
            records["empty"].append({**base, **blank_status})
    return {kind: pd.DataFrame(rows) for kind, rows in records.items()}


def _add_statistics(
    doc: Document,
    df: pd.DataFrame,
    source_schemas: Mapping[str, Iterable[str]] | None,
    cleaning_audit: Iterable[CleaningAuditEntry | Mapping[str, Any]],
) -> None:
    statistics = _business_statistics(df, source_schemas, cleaning_audit)
    labels = {
        "numeric": "Numeric Measures",
        "categorical": "Categorical Columns",
        "date": "Date Columns",
        "identifier": "Identifier Columns",
        "empty": "Empty Columns",
    }
    added = False
    for kind in ("numeric", "categorical", "date", "identifier", "empty"):
        frame = statistics[kind]
        if frame.empty:
            continue
        added = True
        doc.add_heading(labels[kind], level=2)
        _add_table(
            doc,
            list(frame.columns),
            frame.itertuples(index=False, name=None),
        )
    if not added:
        doc.add_paragraph("No column statistics are available for this dataset.")


def _review_rows(outlier_df: pd.DataFrame) -> tuple[list[str], list[list[Any]]]:
    headers = [
        "Record ID",
        "Source File",
        "Product",
        "Field",
        "Value",
        "Upper Review Threshold",
    ]
    rows = []
    for _, record in outlier_df.head(MAX_OUTLIER_ROWS).iterrows():
        upper_value = record.get("Upper Review Boundary")
        if (
            (upper_value is None or pd.isna(upper_value))
            and "Q3" in record
            and "IQR" in record
        ):
            upper_value = record.get("Q3") + 1.5 * record.get("IQR")
        upper = _display(upper_value)
        flagged_value = record.get("Value")
        if isinstance(flagged_value, float):
            flagged_value = f"{flagged_value:.2f}"
        product = record.get("Description", "—")
        if record.get("Review Type") == "Integrity check":
            product = f"Integrity check — {product}"
        rows.append([
            record.get("Record ID", f"Row {record.get('Row', '—')}"),
            record.get("Source File", "—"),
            product,
            friendly_column_name(record.get("Column", "—")),
            flagged_value,
            upper,
        ])
    return headers, rows


def _add_values_to_review(
    doc: Document,
    outlier_df: pd.DataFrame | None,
) -> None:
    if outlier_df is None or outlier_df.empty:
        doc.add_paragraph(
            "No outliers were flagged by the IQR review. "
            "No values currently require statistical follow-up."
        )
        return
    _add_note(
        doc,
        "These values fall outside the statistical review range. "
        "They are not automatically incorrect.",
    )
    for field, field_rows in outlier_df.groupby("Column", dropna=False):
        section_label = (
            "Integrity Checks"
            if field_rows.get("Review Type", pd.Series(dtype=object))
            .eq("Integrity check")
            .all()
            else f"Unusual {friendly_column_name(field)}s"
        )
        doc.add_heading(section_label, level=2)
        headers, rows = _review_rows(field_rows)
        _add_table(
            doc,
            headers,
            rows,
            widths_dxa=[1250, 1600, 1700, 1250, 1500, 2060],
        )
    if len(outlier_df) > MAX_OUTLIER_ROWS:
        doc.add_paragraph(
            f"Showing the first {MAX_OUTLIER_ROWS} of "
            f"{len(outlier_df):,} flagged values."
        )


def _add_charts(doc: Document, chart_images: list[bytes] | None) -> None:
    if not chart_images:
        doc.add_paragraph("No charts were selected for this report.")
        return
    for index, image_bytes in enumerate(chart_images, start=1):
        try:
            doc.add_picture(BytesIO(image_bytes), width=Inches(5.8))
            caption = doc.add_paragraph(f"Chart {index}")
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            logger.warning("Could not embed chart %d", index)
            doc.add_paragraph(f"Chart {index} could not be rendered.")


def _add_preview(doc: Document, df: pd.DataFrame) -> None:
    if df.empty:
        doc.add_paragraph("The final dataset contains no rows to preview.")
        return
    columns = list(df.columns[:PREVIEW_COLUMNS])
    preview = df.loc[:, columns].head(PREVIEW_ROWS)
    _add_table(
        doc,
        [friendly_column_name(column) for column in columns],
        preview.itertuples(index=False, name=None),
    )
    doc.add_paragraph(
        f"Preview only: showing {len(preview):,} of {len(df):,} records "
        f"and {len(columns):,} of {len(df.columns):,} columns."
    )


def _add_methodology(doc: Document) -> None:
    methods_heading = doc.add_heading("Methods Used", level=2)
    methods_heading.paragraph_format.space_before = Pt(6)
    methods_heading.paragraph_format.space_after = Pt(4)
    first_method = doc.add_paragraph(
        "Cells unavailable from a source are identified from the original schema "
        "of each file. Missing values requiring attention are identified only "
        "where the source contained the field. Approved statistical replacements "
        "use valid values from the same source file first, configured business "
        "groups second, and the complete dataset only as a documented fallback."
    )
    first_method.paragraph_format.space_after = Pt(4)
    second_method = doc.add_paragraph(
        "Validated arithmetic relationships are checked after cleaning. When a "
        "missing value can be recovered unambiguously from a trusted relationship, "
        "the formula and inputs are recorded in the Cleaning Audit."
    )
    second_method.paragraph_format.space_after = Pt(4)
    limitations_heading = doc.add_heading("Limitations", level=2)
    limitations_heading.paragraph_format.space_before = Pt(6)
    limitations_heading.paragraph_format.space_after = Pt(4)
    limitations = [
        "Statistical replacements are estimates and should be validated before "
        "operational use.",
        "Unusual values are not automatically errors; they require business review.",
        "Results depend on user-selected cleaning strategies.",
        "The toolkit does not generate business judgments automatically.",
        "Identifier, measure, and category detection is based on column names, "
        "data types, and uniqueness patterns and may need domain-specific review.",
        "Relationship detection is limited to configured rules or relationships "
        "supported by almost all complete records.",
    ]
    for text in limitations:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.1
        run = paragraph.add_run(text)
        _set_run_font(run, size=11, color=NAVY)


def generate_report(
    df: pd.DataFrame,
    stats_df: pd.DataFrame,
    source_files: list[str],
    cleaning_report: dict | None = None,
    outlier_df: pd.DataFrame | None = None,
    chart_images: list[bytes] | None = None,
    include_sections: list[str] | tuple[str, ...] | None = None,
    *,
    cleaning_audit: Iterable[
        CleaningAuditEntry | Mapping[str, Any]
    ] | None = None,
    source_schemas: Mapping[str, Iterable[str]] | None = None,
    dataset_name: str = "Approved cleaned dataset",
    generated_at: datetime | None = None,
    report_id: str | None = None,
) -> bytes:
    """Generate an enterprise-quality Word data quality report.

    ``stats_df`` remains accepted for API compatibility; the report derives its
    type-specific business tables directly from ``df`` to avoid identifier
    averages and sparse, overly wide pandas summary tables.
    """
    del stats_df
    include_sections = list(
        DEFAULT_SECTIONS if include_sections is None else include_sections
    )
    selected = [section for section in SECTION_ORDER if section in include_sections]
    generated_at = generated_at or datetime.now()
    report_id = report_id or (
        f"DQ-{generated_at:%Y%m%d}-{uuid4().hex[:8].upper()}"
    )
    cleaning_report = dict(cleaning_report or {})
    audit = list(cleaning_audit or cleaning_report.get("audit", []))

    if not cleaning_report:
        missing = classify_missing_values(df, source_schemas)
        integrity_report = validate_integrity(df)
        cleaning_report = {
            "rows_before": len(df),
            "rows_after": len(df),
            "duplicates_removed": 0,
            "incomplete_rows_removed": 0,
            "row_level_missing_before": int(missing["row_level_count"].sum()),
            "row_level_missing_after": int(missing["row_level_count"].sum()),
            "structural_missing_before": int(missing["structural_count"].sum()),
            "structural_missing_after": int(missing["structural_count"].sum()),
            "values_filled": 0,
            "estimated_values": 0,
            "deterministic_recoveries": 0,
            "integrity_passed": integrity_report.passed,
            "integrity_issue_count": len(integrity_report.issues),
            "severe_integrity_issue_count": integrity_report.severe_count,
            "integrity_issues": integrity_report.issue_records(),
            "warnings": [],
        }
        integrity_findings = integrity_issues_frame(integrity_report)
        if not integrity_findings.empty:
            outlier_df = pd.concat(
                [
                    outlier_df if outlier_df is not None else pd.DataFrame(),
                    integrity_findings,
                ],
                ignore_index=True,
                sort=False,
            )

    try:
        doc = Document()
        _configure_document(doc)
        _add_running_header(doc, report_id)
        _add_page_footer(doc, generated_at)
        _add_title_area(
            doc,
            generated_at=generated_at,
            dataset_name=dataset_name,
            source_files=source_files,
            report_id=report_id,
        )

        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(
            _executive_summary(
                df,
                source_files,
                cleaning_report,
                outlier_df,
            )
        )

        if int(cleaning_report.get("structural_missing_after", 0)):
            _add_note(
                doc,
                "Some cells remain blank because they came from source files that "
                "did not contain the corresponding field. No action is required "
                "for those cells.",
            )
        if int(cleaning_report.get("estimated_values", 0)):
            _add_note(
                doc,
                "Estimated values are marked in the Cleaning Audit and should be "
                "validated before operational use.",
                warning=True,
            )
        if int(cleaning_report.get("deterministic_recoveries", 0)):
            _add_note(
                doc,
                "Deterministic recoveries were calculated from validated "
                "relationships. Their formulas and inputs are recorded in the "
                "Cleaning Audit.",
            )
        if int(cleaning_report.get("severe_integrity_issue_count", 0)):
            _add_note(
                doc,
                "Relationship validation failed. The dataset must not be described "
                "as business-ready until the severe issues are reviewed.",
                warning=True,
            )
        if (
            int(cleaning_report.get("duplicates_removed", 0))
            or int(cleaning_report.get("incomplete_rows_removed", 0))
        ):
            _add_note(
                doc,
                "Removed rows are listed in the Cleaning Audit.",
                warning=True,
            )

        section_number = 0
        for section_key in selected:
            section_number += 1
            section_heading = doc.add_heading(
                f"{section_number}. {SECTION_TITLES[section_key]}",
                level=1,
            )
            if section_key == "methodology":
                section_heading.paragraph_format.space_before = Pt(8)
                section_heading.paragraph_format.space_after = Pt(4)
            if section_key == "overview":
                _add_overview(doc, df, source_files, source_schemas)
            elif section_key == "quality":
                _add_quality_summary(doc, cleaning_report, audit)
            elif section_key == "statistics":
                _add_statistics(doc, df, source_schemas, audit)
            elif section_key == "outliers":
                _add_values_to_review(doc, outlier_df)
            elif section_key == "charts":
                _add_charts(doc, chart_images)
            elif section_key == "preview":
                _add_preview(doc, df)
            elif section_key == "methodology":
                _add_methodology(doc)

        buffer = BytesIO()
        doc.save(buffer)
        result = buffer.getvalue()
        logger.info(
            "Generated data quality report: %.1f KB, %d numbered sections",
            len(result) / 1024,
            len(selected),
        )
        return result
    except ReportError:
        raise
    except Exception as exc:
        logger.exception("Report generation failed")
        raise ReportError(f"Report generation failed: {exc}") from exc
