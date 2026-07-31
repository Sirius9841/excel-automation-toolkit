"""Generate and validate business-ready sample Excel, CSV, and Word outputs."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
import re
import sys
from datetime import datetime

from docx import Document as DocxDocument
from docx.oxml.ns import qn
from openpyxl import load_workbook
import pandas as pd

PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

from config.settings import OUTPUT_DIR, SAMPLE_DATA_DIR  # noqa: E402
from src.data_processor import merge_datasets, remove_duplicates  # noqa: E402
from src.data_quality import (  # noqa: E402
    apply_missing_value_strategies,
    build_source_schemas,
    classify_missing_values,
    missing_decision_summary,
    structural_blank_audit,
)
from src.integrity import (  # noqa: E402
    detect_validated_relationships,
    validate_integrity,
)
from src.exporter import (  # noqa: E402
    AUDIT_SHEET_NAME,
    EXCEL_SHEET_NAME,
    REVIEW_SHEET_NAME,
    SUMMARY_SHEET_NAME,
    export_audit_to_csv,
    export_to_csv,
    export_to_excel,
)
from src.report_generator import generate_report  # noqa: E402
from src.workflow import DEFAULT_REPORT_SECTIONS, standard_report_context  # noqa: E402


def _document_text(document: DocxDocument) -> str:
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return " ".join(parts)


def _normalized_output_value(value: object) -> str:
    if value is None or bool(pd.isna(value)):
        return ""
    if isinstance(value, (pd.Timestamp, datetime)):
        return pd.Timestamp(value).strftime("%Y-%m-%d")
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value)


def generate_outputs() -> dict[str, Path]:
    """Create the requested output files and return their paths."""
    sources = [
        (path.name, pd.read_excel(path))
        for path in (
            SAMPLE_DATA_DIR / "sales_north.xlsx",
            SAMPLE_DATA_DIR / "sales_south.xlsx",
        )
    ]
    source_schemas = build_source_schemas(sources)
    merged, _ = merge_datasets(sources)
    deduplicated, duplicate_report = remove_duplicates(merged)
    relationships = detect_validated_relationships(deduplicated)
    cleaning_result = apply_missing_value_strategies(
        deduplicated,
        {
            "quantity": "recover_relationship",
            "region": "leave_blank",
            "discount_code": "leave_blank",
            "customer_city": "leave_blank",
        },
        source_schemas,
    )
    cleaned = cleaning_result.cleaned
    integrity_report = validate_integrity(cleaned, relationships)
    assert integrity_report.passed
    cleaning_audit = list(duplicate_report.get("audit", [])) + [
        entry.as_record() for entry in cleaning_result.audit
    ]
    context = standard_report_context(
        cleaned,
        merged,
        duplicate_report,
        list(cleaning_result.messages),
        source_schemas=source_schemas,
        cleaning_audit=cleaning_audit,
    )

    excel_bytes, _ = export_to_excel(
        cleaned,
        "business_ready_cleaned_data",
        cleaning_report=context["cleaning_report"],
        cleaning_audit=cleaning_audit,
        outlier_df=context["outlier_df"],
        source_schemas=source_schemas,
        source_files=[name for name, _ in sources],
    )
    csv_bytes, _ = export_to_csv(cleaned, "business_ready_cleaned_data")
    report_bytes = generate_report(
        cleaned,
        context["stats_df"],
        [name for name, _ in sources],
        cleaning_report=context["cleaning_report"],
        outlier_df=context["outlier_df"],
        include_sections=DEFAULT_REPORT_SECTIONS,
        cleaning_audit=cleaning_audit,
        source_schemas=source_schemas,
        dataset_name="Business-ready cleaned sales data",
    )
    complete_audit = cleaning_audit + [
        entry.as_record()
        for entry in structural_blank_audit(cleaned, source_schemas)
    ]
    audit_csv_bytes, _ = export_audit_to_csv(
        complete_audit,
        "business_ready_cleaning_audit",
    )

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    paths = {
        "excel": OUTPUT_DIR / "business_ready_cleaned_data.xlsx",
        "csv": OUTPUT_DIR / "business_ready_cleaned_data.csv",
        "word": OUTPUT_DIR / "business_ready_data_quality_report.docx",
        "audit_csv": OUTPUT_DIR / "business_ready_cleaning_audit.csv",
    }
    paths["excel"].write_bytes(excel_bytes)
    paths["csv"].write_bytes(csv_bytes)
    paths["word"].write_bytes(report_bytes)
    paths["audit_csv"].write_bytes(audit_csv_bytes)
    _verify_outputs(paths, cleaned, source_schemas)
    return paths


def _verify_outputs(
    paths: dict[str, Path],
    cleaned: pd.DataFrame,
    source_schemas: dict[str, frozenset[str]],
) -> None:
    workbook = load_workbook(paths["excel"], data_only=True)
    assert workbook.sheetnames == [
        EXCEL_SHEET_NAME,
        SUMMARY_SHEET_NAME,
        REVIEW_SHEET_NAME,
        AUDIT_SHEET_NAME,
    ]
    assert "CleanedDataTable" in workbook[EXCEL_SHEET_NAME].tables

    document = DocxDocument(paths["word"])
    full_text = _document_text(document)
    assert not re.search(r"\b(?:nan|NaT|None|fill_mode|fill_median)\b", full_text)
    assert "Relationship validation passed" in full_text
    numbered = [
        int(match.group(1))
        for paragraph in document.paragraphs
        if (match := re.match(r"^(\d+)\.", paragraph.text))
    ]
    assert numbered == list(range(1, len(numbered) + 1))

    csv_frame = pd.read_csv(paths["csv"], dtype=str, keep_default_na=False)
    worksheet = workbook[EXCEL_SHEET_NAME]
    assert len(csv_frame) == worksheet.max_row - 1 == len(cleaned)
    internal_headers = []
    for cell in worksheet[1]:
        assert cell.comment is not None
        internal_headers.append(
            cell.comment.text.split("Internal field name: ", 1)[1]
        )
    excel_values = pd.DataFrame(
        [
            [_normalized_output_value(value) for value in row]
            for row in worksheet.iter_rows(
                min_row=2,
                values_only=True,
            )
        ],
        columns=internal_headers,
    )
    pd.testing.assert_frame_equal(
        excel_values.reset_index(drop=True),
        csv_frame.loc[:, internal_headers].reset_index(drop=True),
    )
    quantity_text = csv_frame.set_index("order_id")["quantity"]
    assert quantity_text.loc["1011"] == "14"
    assert quantity_text.loc["1041"] == "3"
    assert quantity_text.loc["2006"] == "26"
    numeric = cleaned[["quantity", "unit_price", "total"]].apply(
        pd.to_numeric,
        errors="coerce",
    )
    complete = numeric.dropna()
    differences = (
        complete["total"]
        - complete["quantity"] * complete["unit_price"]
    ).abs()
    assert bool((differences <= 0.01).all())
    recovered = cleaned.set_index("order_id")["quantity"]
    assert recovered.loc[1011] == 14
    assert recovered.loc[1041] == 3
    assert recovered.loc[2006] == 26
    assert int(cleaned[["discount_code", "region", "customer_city"]].isna().sum().sum()) == 112

    north = cleaned["source_file"].eq("sales_north.xlsx")
    south = cleaned["source_file"].eq("sales_south.xlsx")
    assert cleaned.loc[north, "discount_code"].isna().all()
    assert cleaned.loc[south, "customer_city"].isna().all()
    assert "discount_code" not in source_schemas["sales_north.xlsx"]
    assert "customer_city" not in source_schemas["sales_south.xlsx"]

    audit_frame = pd.read_csv(paths["audit_csv"], dtype=str, keep_default_na=False)
    missing_summary = classify_missing_values(cleaned, source_schemas)
    assert int(missing_summary["row_level_count"].sum()) == 14
    assert int(missing_summary["structural_count"].sum()) == 98
    decisions = missing_decision_summary(
        missing_summary,
        audit_frame.to_dict("records"),
    )
    assert decisions.approved_unchanged == 14
    assert decisions.pending_review == 0
    assert decisions.failed_or_unresolved == 0
    structural = audit_frame["action_type"].eq("structural_blank")
    assert int(structural.sum()) == 2
    approved = audit_frame["decision_state"].eq("approved_unchanged")
    approved_count = pd.to_numeric(
        audit_frame.loc[approved, "affected_row_count"],
        errors="coerce",
    ).fillna(0).sum()
    assert int(approved_count) == 14
    assert audit_frame["audit_event_id"].nunique() == len(audit_frame)
    removed = pd.to_numeric(
        audit_frame.get("rows_removed", pd.Series(dtype=int)),
        errors="coerce",
    ).fillna(0).sum()
    assert int(removed) == 2
    assert "Seventeen missing values were reviewed" in full_text
    assert "fourteen values were approved to remain blank" in full_text
    assert "14 missing values require attention" not in full_text

    summary_rows = [
        tuple(cell.value for cell in row)
        for row in workbook[SUMMARY_SHEET_NAME].iter_rows()
    ]
    status_header = next(
        index
        for index, row in enumerate(summary_rows)
        if row[0] == "Blank status by field"
    )
    excel_statuses = {
        row[0]: row[1:4]
        for row in summary_rows[status_header + 1:]
        if row[0] in {"Region", "Customer City", "Discount Code"}
    }
    assert excel_statuses == {
        "Region": (2, 0, 0),
        "Customer City": (1, 40, 0),
        "Discount Code": (11, 58, 0),
    }

    category_table = next(
        table
        for table in document.tables
        if "Approved blank" in [
            cell.text for cell in table.rows[0].cells
        ]
        and any(
            row.cells[0].text == "Customer City"
            for row in table.rows[1:]
        )
    )
    category_headers = [
        cell.text for cell in category_table.rows[0].cells
    ]
    word_statuses = {
        row.cells[0].text: {
            header: row.cells[index].text
            for index, header in enumerate(category_headers)
        }
        for row in category_table.rows[1:]
    }
    assert word_statuses["Region"]["Approved blank"] == "2"
    assert word_statuses["Customer City"]["Approved blank"] == "1"
    assert word_statuses["Customer City"]["Unavailable from source"] == "40"
    assert word_statuses["Discount Code"]["Approved blank"] == "11"
    assert word_statuses["Discount Code"]["Unavailable from source"] == "58"
    assert all(
        word_statuses[field]["Decisions pending"] == "0"
        for field in ("Region", "Customer City", "Discount Code")
    )

    source_table = next(
        table
        for table in document.tables
        if [cell.text for cell in table.rows[0].cells]
        == ["Source File", "Records"]
    )
    assert len(source_table.rows) == 3
    assert all(
        row._tr.get_or_add_trPr().find(qn("w:cantSplit")) is not None
        for row in source_table.rows
    )
    assert all(
        paragraph.paragraph_format.keep_with_next
        for row in source_table.rows[:-1]
        for cell in row.cells
        for paragraph in cell.paragraphs
    )


if __name__ == "__main__":
    generated = generate_outputs()
    for label, path in generated.items():
        print(f"{label}: {path}")
