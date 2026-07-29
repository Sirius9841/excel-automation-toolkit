"""Tests for report_generator.py — Word report generation."""

import sys
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

import pandas as pd
import pytest
from docx import Document as DocxDocument
from io import BytesIO

from src.report_generator import (
    generate_report,
    ReportError,
)
from src.analyzer import generate_summary_statistics


@pytest.fixture
def df_sample() -> pd.DataFrame:
    return pd.DataFrame({
        "product": ["Widget A", "Widget B", "Gadget X", "Gadget Y", "Widget C"],
        "price": [10.50, 25.00, 99.99, 150.00, 12.00],
        "quantity": [5, 3, 1, 2, 8],
    })


@pytest.fixture
def stats_df(df_sample) -> pd.DataFrame:
    return generate_summary_statistics(df_sample)


@pytest.fixture
def cleaning_report() -> dict:
    return {
        "rows_before": 100,
        "rows_after": 95,
        "duplicates_removed": 3,
        "total_missing_before": 12,
        "total_missing_after": 2,
        "missing_actions": [
            "Column 'price': filled 5 missing values (fill_median)",
            "Column 'product': filled 2 missing values (fill_mode)",
        ],
    }


@pytest.fixture
def outlier_df() -> pd.DataFrame:
    return pd.DataFrame({
        "Column": ["price", "price"],
        "Row": [3, 7],
        "Value": [999.0, 0.01],
        "Q1": [12.0, 12.0],
        "Q3": [80.0, 80.0],
        "IQR": [68.0, 68.0],
    })


class TestGenerateReport:

    def test_returns_bytes(self, df_sample, stats_df):
        result = generate_report(
            df=df_sample,
            stats_df=stats_df,
            source_files=["test.csv"],
        )
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_valid_docx(self, df_sample, stats_df):
        result = generate_report(
            df=df_sample,
            stats_df=stats_df,
            source_files=["test.csv"],
        )
        doc = DocxDocument(BytesIO(result))
        assert len(doc.paragraphs) > 0

    def test_expected_headings_present(self, df_sample, stats_df):
        result = generate_report(
            df=df_sample,
            stats_df=stats_df,
            source_files=["test.csv"],
        )
        doc = DocxDocument(BytesIO(result))
        heading_texts = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert "1. Dataset Overview" in heading_texts
        assert "2. Data Quality Overview" in heading_texts
        assert "3. Summary Statistics" in heading_texts
        assert "4. Outlier Summary" in heading_texts
        assert "6. Data Preview" in heading_texts
        assert "7. Methodology & Limitations" in heading_texts

    def test_with_cleaning_report(self, df_sample, stats_df, cleaning_report):
        result = generate_report(
            df=df_sample,
            stats_df=stats_df,
            source_files=["test.csv"],
            cleaning_report=cleaning_report,
        )
        doc = DocxDocument(BytesIO(result))
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "3" in full_text  # duplicates removed
        assert "fill_median" in full_text

    def test_with_outliers(self, df_sample, stats_df, outlier_df):
        result = generate_report(
            df=df_sample,
            stats_df=stats_df,
            source_files=["test.csv"],
            outlier_df=outlier_df,
        )
        doc = DocxDocument(BytesIO(result))
        # Check table cells for outlier values
        table_text = ""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    table_text += cell.text + " "
        assert "price" in table_text
        assert "999.0" in table_text

    def test_empty_outlier_df(self, df_sample, stats_df):
        result = generate_report(
            df=df_sample,
            stats_df=stats_df,
            source_files=["test.csv"],
            outlier_df=pd.DataFrame(),
        )
        doc = DocxDocument(BytesIO(result))
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "No outliers were flagged" in full_text

    def test_select_sections_only(self, df_sample, stats_df):
        result = generate_report(
            df=df_sample,
            stats_df=stats_df,
            source_files=["test.csv"],
            include_sections=["overview", "methodology"],
        )
        doc = DocxDocument(BytesIO(result))
        heading_texts = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert "1. Dataset Overview" in heading_texts
        assert "7. Methodology & Limitations" in heading_texts
        assert "3. Summary Statistics" not in heading_texts

    def test_empty_dataframe(self, stats_df):
        result = generate_report(
            df=pd.DataFrame(),
            stats_df=stats_df,
            source_files=["test.csv"],
        )
        assert isinstance(result, bytes)

    def test_dataframe_not_modified(self, df_sample, stats_df):
        original = df_sample.copy()
        generate_report(
            df=df_sample,
            stats_df=stats_df,
            source_files=["test.csv"],
        )
        pd.testing.assert_frame_equal(df_sample, original)
