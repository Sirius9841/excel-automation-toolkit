"""Regression tests for business-safe cleaning and enterprise outputs."""

from io import BytesIO
from pathlib import Path
import re

from docx import Document as DocxDocument
from docx.oxml.ns import qn
from openpyxl import load_workbook
import pandas as pd

from config.settings import SAMPLE_DATA_DIR
from src.analyzer import (
    generate_summary_statistics,
    summarize_outliers,
)
from src.data_processor import (
    generate_cleaning_report,
    merge_datasets,
    recommend_missing_strategy,
    remove_duplicates,
)
from src.data_quality import (
    apply_missing_value_strategies,
    audit_summary,
    build_review_summary,
    build_source_schemas,
    classify_missing_values,
    cleaning_change_preview,
    grouped_audit_change_table,
    grouped_audit_change_summaries,
    missing_completion_counts,
    missing_decision_summary,
    missing_status_by_column,
    recommendation_explanation,
    structural_blank_audit,
)
from src.exporter import (
    AUDIT_SHEET_NAME,
    EXCEL_SHEET_NAME,
    REVIEW_SHEET_NAME,
    SUMMARY_SHEET_NAME,
    export_to_csv,
    export_to_excel,
)
from src.file_handler import _identifier_dtype_map
from src.report_generator import generate_report
from src.integrity import (
    combine_review_findings,
    detect_validated_relationships,
    validate_integrity,
)
from src.workflow import standard_report_context


def _all_docx_text(data: bytes) -> str:
    document = DocxDocument(BytesIO(data))
    text = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            text.extend(cell.text for cell in row.cells)
    return " ".join(text)


def _schema_fixture() -> tuple[
    list[tuple[str, pd.DataFrame]],
    dict[str, frozenset[str]],
    pd.DataFrame,
]:
    north = pd.DataFrame({
        "order_id": ["001", "002", "003"],
        "quantity": [1.0, None, 5.0],
        "customer_city": ["Lyon", "Paris", None],
    })
    south = pd.DataFrame({
        "order_id": ["004", "005"],
        "quantity": [10.0, 20.0],
        "discount_code": ["DISC10", None],
    })
    sources = [("north.xlsx", north), ("south.xlsx", south)]
    schemas = build_source_schemas(sources)
    merged, _ = merge_datasets(sources)
    return sources, schemas, merged


def test_absent_source_column_is_structural_missingness():
    _, schemas, merged = _schema_fixture()
    summary = classify_missing_values(merged, schemas).set_index("column")
    assert summary.loc["discount_code", "structural_count"] == 3
    assert summary.loc["customer_city", "structural_count"] == 2
    assert summary.loc["discount_code", "row_level_count"] == 1


def test_structural_blanks_are_not_filled_by_default():
    _, schemas, merged = _schema_fixture()
    result = apply_missing_value_strategies(
        merged,
        {"discount_code": "fill_mode", "customer_city": "fill_mode"},
        schemas,
    )
    north = result.cleaned["source_file"].eq("north.xlsx")
    south = result.cleaned["source_file"].eq("south.xlsx")
    assert result.cleaned.loc[north, "discount_code"].isna().all()
    assert result.cleaned.loc[south, "customer_city"].isna().all()


def test_structural_blanks_are_not_selected_for_row_removal():
    _, schemas, merged = _schema_fixture()
    result = apply_missing_value_strategies(
        merged,
        {"discount_code": "drop_rows"},
        schemas,
    )
    assert len(result.cleaned) == len(merged) - 1
    assert result.cleaned["source_file"].eq("north.xlsx").sum() == 3


def test_structural_only_column_recommendation_is_leave_blank():
    first = pd.DataFrame({"order_id": [1], "discount_code": ["DISC10"]})
    second = pd.DataFrame({"order_id": [2]})
    sources = [("with_code.xlsx", first), ("without_code.xlsx", second)]
    schemas = build_source_schemas(sources)
    merged, _ = merge_datasets(sources)
    assert (
        recommend_missing_strategy(merged, "discount_code", schemas)
        == "leave_blank"
    )


def test_existing_source_column_blank_is_row_level_missingness():
    _, schemas, merged = _schema_fixture()
    summary = classify_missing_values(merged, schemas).set_index("column")
    assert summary.loc["quantity", "structural_count"] == 0
    assert summary.loc["quantity", "row_level_count"] == 1


def test_group_aware_median_uses_same_source_file():
    north = pd.DataFrame({"value": [1.0, None, 101.0]})
    south = pd.DataFrame({"value": [1000.0, 2000.0, 3000.0]})
    sources = [("north.xlsx", north), ("south.xlsx", south)]
    schemas = build_source_schemas(sources)
    merged, _ = merge_datasets(sources)
    result = apply_missing_value_strategies(
        merged,
        {"value": "fill_median"},
        schemas,
    )
    filled = result.cleaned.loc[
        result.cleaned["source_file"].eq("north.xlsx"),
        "value",
    ]
    assert 51.0 in filled.tolist()
    assert result.audit[0].strategy_scope == "Within source file: north.xlsx"


def test_global_fallback_is_documented():
    first = pd.DataFrame({"value": [None, None]})
    second = pd.DataFrame({"value": [10.0, 30.0]})
    sources = [("empty.xlsx", first), ("values.xlsx", second)]
    schemas = build_source_schemas(sources)
    merged, _ = merge_datasets(sources)
    result = apply_missing_value_strategies(
        merged,
        {"value": "fill_median"},
        schemas,
    )
    assert result.cleaned.loc[
        result.cleaned["source_file"].eq("empty.xlsx"),
        "value",
    ].eq(20.0).all()
    assert {
        entry.strategy_scope for entry in result.audit
    } == {"Global fallback"}


def test_original_frames_are_never_modified():
    sources, schemas, merged = _schema_fixture()
    originals = [(name, frame.copy(deep=True)) for name, frame in sources]
    merged_original = merged.copy(deep=True)
    apply_missing_value_strategies(
        merged,
        {"quantity": "fill_median", "discount_code": "drop_rows"},
        schemas,
    )
    pd.testing.assert_frame_equal(merged, merged_original)
    for (_, frame), (_, original) in zip(sources, originals):
        pd.testing.assert_frame_equal(frame, original)


def test_cleaning_audit_records_replacements_and_removals():
    _, schemas, merged = _schema_fixture()
    result = apply_missing_value_strategies(
        merged,
        {"quantity": "fill_median", "discount_code": "drop_rows"},
        schemas,
    )
    summary = audit_summary(result.audit)
    assert summary["values_filled"] == 1
    assert summary["incomplete_rows_removed"] == 1
    assert all(entry.row_identifier for entry in result.audit)
    assert all(entry.source_file for entry in result.audit)


def _enterprise_export_fixture() -> tuple[
    pd.DataFrame,
    dict[str, frozenset[str]],
    list[dict],
    dict,
    pd.DataFrame,
]:
    _, schemas, merged = _schema_fixture()
    result = apply_missing_value_strategies(
        merged,
        {"quantity": "fill_median", "discount_code": "fill_mode"},
        schemas,
    )
    outliers = pd.DataFrame([{
        "Record ID": "003",
        "Source File": "north.xlsx",
        "Description": "—",
        "Column": "quantity",
        "Value": 500,
        "Lower Review Boundary": -10,
        "Upper Review Boundary": 100,
        "Reason": "Outside review range.",
    }])
    report = generate_cleaning_report(
        merged,
        result.cleaned,
        source_schemas=schemas,
        cleaning_audit=result.audit,
        missing_actions=list(result.messages),
    )
    return (
        result.cleaned,
        schemas,
        [entry.as_record() for entry in result.audit],
        report,
        outliers,
    )


def test_excel_contains_required_sheets_and_filterable_table():
    df, schemas, audit, report, outliers = _enterprise_export_fixture()
    data, _ = export_to_excel(
        df,
        "enterprise",
        cleaning_report=report,
        cleaning_audit=audit,
        outlier_df=outliers,
        source_schemas=schemas,
        source_files=list(schemas),
    )
    workbook = load_workbook(BytesIO(data))
    assert workbook.sheetnames == [
        EXCEL_SHEET_NAME,
        SUMMARY_SHEET_NAME,
        REVIEW_SHEET_NAME,
        AUDIT_SHEET_NAME,
    ]
    cleaned = workbook[EXCEL_SHEET_NAME]
    assert "CleanedDataTable" in cleaned.tables
    assert cleaned.freeze_panes == "A2"
    assert cleaned.tables["CleanedDataTable"].autoFilter is not None


def test_excel_dates_numbers_and_identifiers_are_formatted():
    df = pd.DataFrame({
        "order_id": ["0001", "0002"],
        "order_date": pd.to_datetime(["2025-01-01", "2025-02-03"]),
        "quantity": [6.0, 7.0],
        "unit_price": [12.5, 20.0],
        "source_file": ["a.xlsx", "a.xlsx"],
    })
    schemas = build_source_schemas([("a.xlsx", df.drop(columns="source_file"))])
    data, _ = export_to_excel(df, "formats", source_schemas=schemas)
    workbook = load_workbook(BytesIO(data))
    sheet = workbook[EXCEL_SHEET_NAME]
    assert sheet["A2"].value == "0001"
    assert sheet["A2"].number_format == "@"
    assert sheet["B2"].number_format == "yyyy-mm-dd"
    assert sheet["C2"].number_format == "#,##0"
    assert sheet["D2"].number_format == "#,##0.00"


def test_identifier_dtype_map_preserves_id_like_source_columns():
    mapping = _identifier_dtype_map(
        ["order_id", "employee_id", "discount_code", "quantity"]
    )
    assert mapping == {"order_id": "string", "employee_id": "string"}


def test_csv_matches_cleaned_data_and_preserves_non_english_text():
    df = pd.DataFrame({
        "order_id": ["0001", "0002"],
        "customer_city": ["München", "São Paulo"],
        "quantity": [6.0, 7.0],
        "order_date": pd.to_datetime(["2025-01-01", "2025-02-03"]),
    })
    excel_data, _ = export_to_excel(df, "match")
    csv_data, _ = export_to_csv(df, "match")
    workbook = load_workbook(BytesIO(excel_data), data_only=True)
    sheet = workbook[EXCEL_SHEET_NAME]
    excel_rows = list(sheet.iter_rows(values_only=True))
    csv_frame = pd.read_csv(BytesIO(csv_data), dtype=str, keep_default_na=False)
    assert csv_data.startswith(b"\xef\xbb\xbf")
    assert csv_frame.loc[0, "order_id"] == "0001"
    assert csv_frame.loc[0, "customer_city"] == "München"
    assert str(excel_rows[1][0]) == csv_frame.loc[0, "order_id"]
    assert excel_rows[1][1] == csv_frame.loc[0, "customer_city"]
    assert str(excel_rows[1][2]) == csv_frame.loc[0, "quantity"]
    assert excel_rows[1][3].strftime("%Y-%m-%d") == csv_frame.loc[0, "order_date"]


def test_word_report_has_continuous_sections_and_no_internal_values():
    df, schemas, audit, report, outliers = _enterprise_export_fixture()
    data = generate_report(
        df,
        generate_summary_statistics(df),
        list(schemas),
        cleaning_report=report,
        outlier_df=outliers,
        include_sections=[
            "overview",
            "quality",
            "statistics",
            "outliers",
            "methodology",
        ],
        cleaning_audit=audit,
        source_schemas=schemas,
    )
    document = DocxDocument(BytesIO(data))
    headings = [
        paragraph.text
        for paragraph in document.paragraphs
        if paragraph.style.name == "Heading 1"
        and re.match(r"^\d+\.", paragraph.text)
    ]
    assert [int(text.split(".", 1)[0]) for text in headings] == [1, 2, 3, 4, 5]
    full_text = _all_docx_text(data)
    assert not re.search(r"\b(?:nan|NaT|None)\b", full_text)
    assert "fill_mode" not in full_text
    assert "fill_median" not in full_text
    assert "source files that did not contain" in full_text


def test_identifier_statistics_do_not_show_averages():
    df = pd.DataFrame({
        "order_id": [1001, 1002, 1003, 1004],
        "quantity": [1, 2, 3, 4],
    })
    data = generate_report(
        df,
        generate_summary_statistics(df),
        ["source.xlsx"],
        include_sections=["statistics"],
    )
    document = DocxDocument(BytesIO(data))
    identifier_table = next(
        table
        for table in document.tables
        if table.rows[0].cells[0].text == "Column"
        and any(
            cell.text == "Repeated Identifiers"
            for cell in table.rows[0].cells
        )
    )
    headers = [cell.text for cell in identifier_table.rows[0].cells]
    assert "Average" not in headers
    assert "Median" not in headers


def test_outlier_context_includes_record_identifier_and_source():
    df = pd.DataFrame({
        "order_id": list(range(1, 11)),
        "product": [f"Product {value}" for value in range(1, 11)],
        "source_file": ["source.xlsx"] * 10,
        "total": [10, 11, 12, 10, 11, 12, 10, 11, 12, 1000],
    })
    outliers = summarize_outliers(df, ["total"])
    assert outliers.iloc[0]["Record ID"] == 10
    assert outliers.iloc[0]["Source File"] == "source.xlsx"
    assert outliers.iloc[0]["Description"] == "Product 10"
    assert "Lower Review Boundary" in outliers.columns


def test_reports_reopen_and_zero_row_result_is_handled():
    empty = pd.DataFrame(columns=["order_id", "quantity"])
    report = generate_report(
        empty,
        generate_summary_statistics(empty),
        ["empty.xlsx"],
    )
    document = DocxDocument(BytesIO(report))
    assert document.paragraphs


def test_existing_sample_files_preserve_structural_blanks():
    sources = [
        (path.name, pd.read_excel(path))
        for path in (
            SAMPLE_DATA_DIR / "sales_north.xlsx",
            SAMPLE_DATA_DIR / "sales_south.xlsx",
        )
    ]
    schemas = build_source_schemas(sources)
    merged, _ = merge_datasets(sources)
    result = apply_missing_value_strategies(
        merged,
        {
            "quantity": "fill_median",
            "region": "fill_mode",
            "customer_city": "fill_mode",
            "discount_code": "fill_mode",
        },
        schemas,
    )
    north = result.cleaned["source_file"].eq("sales_north.xlsx")
    south = result.cleaned["source_file"].eq("sales_south.xlsx")
    assert result.cleaned.loc[north, "discount_code"].isna().all()
    assert result.cleaned.loc[south, "customer_city"].isna().all()
    assert not result.cleaned.loc[north, "discount_code"].eq("DISC10").any()
    assert not result.cleaned.loc[south, "customer_city"].eq("München").any()
    assert result.cleaned["quantity"].isna().sum() == 0
    assert result.cleaned["region"].isna().sum() == 0
    summary = classify_missing_values(result.cleaned, schemas).set_index("column")
    assert summary.loc["discount_code", "structural_count"] == len(
        sources[0][1]
    )
    assert summary.loc["customer_city", "structural_count"] == len(
        sources[1][1]
    )


def _sample_sales_cleaning():
    sources = [
        (path.name, pd.read_excel(path))
        for path in (
            SAMPLE_DATA_DIR / "sales_north.xlsx",
            SAMPLE_DATA_DIR / "sales_south.xlsx",
        )
    ]
    schemas = build_source_schemas(sources)
    merged, _ = merge_datasets(sources)
    deduplicated, duplicate_report = remove_duplicates(merged)
    relationships = detect_validated_relationships(deduplicated)
    result = apply_missing_value_strategies(
        deduplicated,
        {
            "quantity": "recover_relationship",
            "region": "leave_blank",
            "customer_city": "leave_blank",
            "discount_code": "leave_blank",
        },
        schemas,
    )
    audit = list(duplicate_report["audit"]) + [
        entry.as_record() for entry in result.audit
    ]
    return sources, schemas, merged, result, relationships, audit


def test_sample_relationship_is_validated_on_complete_rows():
    _, _, merged, _, relationships, _ = _sample_sales_cleaning()
    relationship = relationships[0]
    assert relationship.expected_relationship == (
        "Total = Quantity × Unit Price"
    )
    assert relationship.complete_rows >= 90
    assert relationship.pass_rate >= 0.98
    complete = merged[["quantity", "unit_price", "total"]].dropna()
    assert bool(
        (
            complete["total"]
            - complete["quantity"] * complete["unit_price"]
        ).abs().le(0.01).all()
    )


def test_missing_quantities_use_deterministic_recovery():
    _, _, _, result, _, _ = _sample_sales_cleaning()
    quantities = result.cleaned.set_index("order_id")["quantity"]
    assert quantities.loc[1011] == 14
    assert quantities.loc[1041] == 3
    assert quantities.loc[2006] == 26
    quantity_audit = [
        entry
        for entry in result.audit
        if entry.column == "quantity"
    ]
    assert len(quantity_audit) == 3
    assert {
        entry.action_type for entry in quantity_audit
    } == {"deterministic_recovery"}
    assert all("÷" in entry.formula_or_strategy for entry in quantity_audit)
    assert not any(
        "median" in entry.strategy.lower()
        for entry in quantity_audit
    )


def test_median_request_does_not_override_deterministic_recovery():
    sources = [
        (path.name, pd.read_excel(path))
        for path in (
            SAMPLE_DATA_DIR / "sales_north.xlsx",
            SAMPLE_DATA_DIR / "sales_south.xlsx",
        )
    ]
    schemas = build_source_schemas(sources)
    merged, _ = merge_datasets(sources)
    result = apply_missing_value_strategies(
        merged,
        {"quantity": "fill_median"},
        schemas,
    )
    quantities = result.cleaned.set_index("order_id")["quantity"]
    assert [quantities.loc[value] for value in (1011, 1041, 2006)] == [
        14,
        3,
        26,
    ]
    assert all(
        entry.action_type == "deterministic_recovery"
        for entry in result.audit
    )


def test_cleaned_sample_has_no_relationship_violations():
    _, _, _, result, relationships, _ = _sample_sales_cleaning()
    report = validate_integrity(result.cleaned, relationships)
    assert report.passed
    assert report.issues == ()
    csv_data, _ = export_to_csv(result.cleaned, "relationship_clean")
    exported = pd.read_csv(BytesIO(csv_data))
    exported_report = validate_integrity(exported, relationships)
    assert exported_report.passed


def test_business_categories_default_to_leave_blank():
    frame = pd.DataFrame({
        "order_id": [1, 2, 3],
        "discount_code": ["DISC10", None, "DISC10"],
        "customer_city": ["Paris", None, "Paris"],
        "region": ["North", None, "North"],
        "status": ["Open", None, "Open"],
    })
    for column in ("discount_code", "customer_city", "region", "status"):
        assert recommend_missing_strategy(frame, column) == "leave_blank"


def test_severe_integrity_issues_reach_excel_and_word_outputs():
    rows = 100
    frame = pd.DataFrame({
        "order_id": range(1, rows + 1),
        "quantity": [2] * rows,
        "unit_price": [10.0] * rows,
        "total": [20.0] * (rows - 1) + [999.0],
        "source_file": ["source.xlsx"] * rows,
    })
    relationships = detect_validated_relationships(frame)
    integrity = validate_integrity(frame, relationships)
    assert integrity.severe_count == 1
    review = combine_review_findings(pd.DataFrame(), integrity)
    cleaning_report = {
        "rows_before": rows,
        "rows_after": rows,
        "duplicates_removed": 0,
        "row_level_missing_before": 0,
        "row_level_missing_after": 0,
        "structural_missing_before": 0,
        "structural_missing_after": 0,
        "values_filled": 0,
        "estimated_values": 0,
        "deterministic_recoveries": 0,
        "integrity_passed": False,
        "integrity_issue_count": 1,
        "severe_integrity_issue_count": 1,
        "integrity_issues": integrity.issue_records(),
        "warnings": ["One severe relationship issue requires review."],
    }
    excel, _ = export_to_excel(
        frame,
        "integrity",
        cleaning_report=cleaning_report,
        outlier_df=review,
    )
    workbook = load_workbook(BytesIO(excel))
    review_text = " ".join(
        str(cell.value)
        for row in workbook[REVIEW_SHEET_NAME].iter_rows()
        for cell in row
        if cell.value is not None
    )
    assert "Integrity check" in review_text
    assert "Relationship check failed" in review_text
    assert workbook[EXCEL_SHEET_NAME].sheet_properties.tabColor is not None

    word = generate_report(
        frame,
        generate_summary_statistics(frame),
        ["source.xlsx"],
        cleaning_report=cleaning_report,
        outlier_df=review,
    )
    word_text = _all_docx_text(word)
    assert "Relationship validation failed" in word_text
    assert "Integrity check" in word_text
    assert "business-ready" in word_text


def test_integer_like_csv_quantities_do_not_use_decimal_suffix():
    frame = pd.DataFrame({
        "quantity": [6.0, 7.5],
        "unit_price": [10.0, 12.25],
    })
    data, _ = export_to_csv(frame, "integer_like")
    lines = data.decode("utf-8-sig").splitlines()
    assert lines[1].split(",")[0] == "6"
    assert lines[2].split(",")[0] == "7.5"


def test_duplicate_audit_has_unique_event_and_source_row_references():
    sources = [
        ("source.xlsx", pd.DataFrame({
            "order_id": [1, 1, 1],
            "value": [10, 10, 10],
        }))
    ]
    merged, _ = merge_datasets(sources)
    _, report = remove_duplicates(merged)
    assert len(report["audit"]) == 2
    assert len({
        record["audit_event_id"]
        for record in report["audit"]
    }) == 2
    assert len({
        record["original_source_row"]
        for record in report["audit"]
    }) == 2


def test_structural_audit_is_aggregated_by_source_and_column():
    _, schemas, merged = _schema_fixture()
    audit = structural_blank_audit(merged, schemas)
    assert len(audit) == 2
    assert all(entry.row_index == "Aggregated" for entry in audit)
    assert {
        (entry.source_file, entry.column)
        for entry in audit
    } == {
        ("north.xlsx", "discount_code"),
        ("south.xlsx", "customer_city"),
    }


def test_completed_missing_values_ui_copy_is_present():
    source = (Path(__file__).resolve().parent.parent / "app.py").read_text(
        encoding="utf-8"
    )
    assert "Missing-value review complete" in source
    assert "All missing-value decisions have been reviewed." in source
    assert "Values recovered or filled" in source
    assert "Approved to remain blank" in source
    assert "Unavailable from source" in source
    assert "Decisions pending" in source
    assert "What changed" in source
    assert "View reviewed records" in source
    assert "Fields unavailable from their source files" in source
    assert "Intentionally unavailable" in source
    assert "Why are these fields unavailable?" in source
    assert "Continue to review changes" in source
    assert "Recommended:** Leave blank" not in source


def test_missing_change_preview_shows_original_and_current_values():
    audit = [{
        "action_type": "deterministic_recovery",
        "action": "Recovered from Total ÷ Unit Price",
        "column": "quantity",
        "source_file": "north.xlsx",
        "business_record_identifier": "1002",
        "row_identifier": "1002",
        "original_value": None,
        "resulting_value": 4,
        "rows_removed": 0,
    }]

    preview = cleaning_change_preview(audit)
    summaries = grouped_audit_change_summaries(audit)

    assert preview.iloc[0].to_dict() == {
        "Record ID": "1002",
        "Field": "quantity",
        "Original": "Blank",
        "Current": 4,
        "Decision or method": "Recovered from Total ÷ Unit Price",
        "Source file": "north.xlsx",
    }
    assert summaries == [
        "Quantity: 1 value recovered from Total ÷ Unit Price."
    ]
    assert grouped_audit_change_table(audit).iloc[0].to_dict() == {
        "Field": "Quantity",
        "Decision or method": "Recovered from Total ÷ Unit Price",
        "Records affected": 1,
    }


def test_missing_completion_summary_uses_computed_counts():
    missing_summary = pd.DataFrame({
        "structural_count": [58, 40],
        "row_level_count": [0, 0],
    })
    audit = [
        {
            "action_type": "deterministic_recovery",
            "rows_removed": 0,
        }
        for _ in range(17)
    ]

    assert missing_completion_counts(missing_summary, audit) == {
        "values_handled": 17,
        "approved_unchanged": 0,
        "unavailable_from_source": 98,
        "decisions_pending": 0,
        "failed_or_unresolved": 0,
    }


def test_leave_blank_decision_states_and_sample_counts():
    sources = [
        (path.name, pd.read_excel(path))
        for path in (
            SAMPLE_DATA_DIR / "sales_north.xlsx",
            SAMPLE_DATA_DIR / "sales_south.xlsx",
        )
    ]
    schemas = build_source_schemas(sources)
    merged, _ = merge_datasets(sources)
    deduplicated, duplicate_report = remove_duplicates(merged)
    before = classify_missing_values(deduplicated, schemas)
    assert int(before["row_level_count"].sum()) == 17
    assert missing_decision_summary(before, []).pending_review == 17

    result = apply_missing_value_strategies(
        deduplicated,
        {
            "quantity": "recover_relationship",
            "discount_code": "leave_blank",
            "region": "leave_blank",
            "customer_city": "leave_blank",
        },
        schemas,
        recorded_at="2026-07-30T12:00:00+00:00",
    )
    after = classify_missing_values(result.cleaned, schemas)
    summary = result.decision_summary
    assert summary.changed == 3
    assert summary.approved_unchanged == 14
    assert summary.unavailable_from_source == 98
    assert summary.pending_review == 0
    assert summary.failed_or_unresolved == 0
    assert summary.complete
    assert int(after["row_level_count"].sum()) == 14
    assert int(result.cleaned.isna().sum().sum()) == 112

    approved = [
        entry for entry in result.audit
        if entry.decision_state == "approved_unchanged"
        or getattr(entry.decision_state, "value", "") == "approved_unchanged"
    ]
    assert sum(entry.affected_row_count for entry in approved) == 14
    assert all(entry.action == "Approved to remain blank" for entry in approved)
    assert all(entry.resulting_state == "Blank retained by user decision" for entry in approved)
    assert int(duplicate_report["removed"]) == 2


def test_leave_blank_export_and_report_keep_blanks_but_not_pending():
    sources, schemas, merged, result, _, audit = _sample_sales_cleaning()
    duplicate_report = {"removed": 2, "audit": audit[:2]}
    report = generate_cleaning_report(
        merged,
        result.cleaned,
        duplicate_report,
        list(result.messages),
        source_schemas=schemas,
        cleaning_audit=audit,
        integrity_report=result.integrity_report,
    )
    assert report["missing_values_reviewed"] == 17
    assert report["values_changed"] == 3
    assert report["approved_unchanged"] == 14
    assert report["decisions_pending"] == 0
    assert report["unavailable_from_source"] == 98
    assert report["integrity_failures"] == 0

    csv_bytes, _ = export_to_csv(result.cleaned, "approved_blanks")
    csv_frame = pd.read_csv(BytesIO(csv_bytes), keep_default_na=False)
    assert int(csv_frame.eq("").sum().sum()) == 112

    report_bytes = generate_report(
        result.cleaned,
        generate_summary_statistics(result.cleaned),
        [name for name, _ in sources],
        cleaning_report=report,
        outlier_df=summarize_outliers(
            result.cleaned,
            ["quantity", "unit_price", "total"],
        ),
        cleaning_audit=audit,
        source_schemas=schemas,
    )
    document = DocxDocument(BytesIO(report_bytes))
    text = " ".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        ]
    )
    assert "Seventeen missing values were reviewed" in text
    assert "fourteen values were approved to remain blank" in text
    assert "14 missing values require attention" not in text


def test_completion_summary_never_derives_attention_from_physical_nulls():
    _, schemas, _, result, _, audit = _sample_sales_cleaning()
    missing = classify_missing_values(result.cleaned, schemas)
    summary = build_review_summary(
        missing,
        audit,
        duplicate_rows_removed=2,
    )

    assert int(result.cleaned.isna().sum().sum()) == 112
    assert summary.recovered_or_filled_count == 3
    assert summary.approved_unchanged_count == 14
    assert summary.unavailable_from_source_count == 98
    assert summary.pending_decision_count == 0
    assert summary.integrity_status == "Passed"
    completion_text = " ".join(summary.activity_lines())
    assert "14 missing values require attention" not in completion_text
    assert "0 decisions pending" in completion_text
    assert "Approved 14 values to remain blank" in completion_text


def test_recommendation_copy_explains_safe_defaults_and_advanced_mode():
    leave_title, leave_detail = recommendation_explanation("leave_blank")
    recover_title, recover_detail = recommendation_explanation(
        "recover_relationship"
    )

    assert leave_title == "Recommended: Leave blank"
    assert "unverified information" in leave_detail
    assert recover_title == (
        "Recommended: Recover from Total ÷ Unit Price"
    )
    assert "without estimation" in recover_detail
    app_source = (
        Path(__file__).resolve().parent.parent / "app.py"
    ).read_text(encoding="utf-8")
    assert "This inserts the most frequent value and may " in app_source
    assert "create incorrect business information." in app_source


def test_schema_compatibility_does_not_change_category_recommendations():
    frame = pd.DataFrame({
        "region": ["North", None, "South"],
        "source_file": ["a.xlsx", "a.xlsx", "a.xlsx"],
        "other": [1, 2, 3],
    })
    highly_similar = {
        "a.xlsx": frozenset({"region", "other", "extra_a", "extra_b"}),
    }
    minimally_similar = {
        "a.xlsx": frozenset({"region"}),
    }

    assert recommend_missing_strategy(
        frame, "region", highly_similar
    ) == "leave_blank"
    assert recommend_missing_strategy(
        frame, "region", minimally_similar
    ) == "leave_blank"


def test_sample_blank_statuses_are_mutually_exclusive_by_field():
    _, schemas, _, result, _, audit = _sample_sales_cleaning()
    statuses = missing_status_by_column(
        result.cleaned,
        schemas,
        audit,
    )

    assert (
        statuses["region"].approved_blank,
        statuses["region"].unavailable_from_source,
        statuses["region"].decisions_pending,
    ) == (2, 0, 0)
    assert (
        statuses["customer_city"].approved_blank,
        statuses["customer_city"].unavailable_from_source,
        statuses["customer_city"].decisions_pending,
    ) == (1, 40, 0)
    assert (
        statuses["discount_code"].approved_blank,
        statuses["discount_code"].unavailable_from_source,
        statuses["discount_code"].decisions_pending,
    ) == (11, 58, 0)


def test_excel_summary_separates_blank_statuses():
    sources, schemas, merged, result, _, audit = _sample_sales_cleaning()
    report = generate_cleaning_report(
        merged,
        result.cleaned,
        {"removed": 2, "audit": audit[:2]},
        list(result.messages),
        source_schemas=schemas,
        cleaning_audit=audit,
        integrity_report=result.integrity_report,
    )
    data, _ = export_to_excel(
        result.cleaned,
        "blank_statuses",
        cleaning_report=report,
        cleaning_audit=audit,
        source_schemas=schemas,
        source_files=[name for name, _ in sources],
    )
    sheet = load_workbook(BytesIO(data), data_only=True)[SUMMARY_SHEET_NAME]
    rows = [tuple(cell.value for cell in row) for row in sheet.iter_rows()]
    header_index = next(
        index
        for index, row in enumerate(rows)
        if row[0] == "Blank status by field"
    )
    status_rows = {
        row[0]: row[1:4]
        for row in rows[header_index + 1:]
        if row[0] in {"Region", "Customer City", "Discount Code"}
    }
    assert status_rows == {
        "Region": (2, 0, 0),
        "Customer City": (1, 40, 0),
        "Discount Code": (11, 58, 0),
    }


def test_word_statistics_separate_blank_statuses_and_keep_source_table():
    sources, schemas, merged, result, _, audit = _sample_sales_cleaning()
    report = generate_cleaning_report(
        merged,
        result.cleaned,
        {"removed": 2, "audit": audit[:2]},
        list(result.messages),
        source_schemas=schemas,
        cleaning_audit=audit,
        integrity_report=result.integrity_report,
    )
    data = generate_report(
        result.cleaned,
        generate_summary_statistics(result.cleaned),
        [name for name, _ in sources],
        cleaning_report=report,
        cleaning_audit=audit,
        source_schemas=schemas,
    )
    document = DocxDocument(BytesIO(data))
    category_table = next(
        table
        for table in document.tables
        if "Approved blank" in [
            cell.text for cell in table.rows[0].cells
        ]
        and any(
            cell.text == "Customer City"
            for row in table.rows[1:]
            for cell in row.cells[:1]
        )
    )
    headers = [cell.text for cell in category_table.rows[0].cells]
    rows = {
        row.cells[0].text: {
            header: row.cells[index].text
            for index, header in enumerate(headers)
        }
        for row in category_table.rows[1:]
    }
    assert rows["Customer City"]["Approved blank"] == "1"
    assert rows["Customer City"]["Unavailable from source"] == "40"
    assert rows["Customer City"]["Decisions pending"] == "0"
    assert rows["Discount Code"]["Approved blank"] == "11"
    assert rows["Discount Code"]["Unavailable from source"] == "58"
    assert rows["Region"]["Approved blank"] == "2"

    source_table = next(
        table
        for table in document.tables
        if [cell.text for cell in table.rows[0].cells]
        == ["Source File", "Records"]
    )
    assert len(source_table.rows) == 3
    for row in source_table.rows:
        assert row._tr.get_or_add_trPr().find(qn("w:cantSplit")) is not None
    for row in source_table.rows[:-1]:
        assert all(
            paragraph.paragraph_format.keep_with_next
            for cell in row.cells
            for paragraph in cell.paragraphs
        )
